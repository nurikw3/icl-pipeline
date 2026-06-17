from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT_PATH = Path("docs/graph_icl_pipeline_report.docx")


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x55, 0x65, 0x75)
LIGHT_FILL = "F2F4F7"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D7DE", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths=None):
    table.autofit = False
    table.style = "Table Grid"
    set_table_width(table)
    set_table_borders(table)
    if widths:
        set_col_widths(table, widths)

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)
        if row_idx == 0:
            set_repeat_table_header(row)
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = INK


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)

    style_table(table, widths=widths)
    document.add_paragraph()
    return table


def add_formula(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    return paragraph


def add_note(document, title, body):
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    set_table_width(table)
    set_table_borders(table, color="D8DEE6", size="4")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.size = Pt(10)
    document.add_paragraph()


def setup_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(document):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("ICL retrieval report: методы, prompts и результаты")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(
        "Краткая техническая статья о pipeline, retrieval, DIPMT prompts, "
        "graph PPR и итоговых метриках"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    metadata = [
        ("Проект", "ICL pipeline для Chagatai -> English"),
        ("Эксперимент", "10 percent smoke tests + 100 percent checks"),
        ("Модель", "openai/gpt-oss-120b"),
        ("BERTScore model", "bert-base-multilingual-cased"),
        ("Дата", date.today().isoformat()),
    ]
    add_table(
        document,
        ["Поле", "Значение"],
        metadata,
        widths=[1.35, 5.05],
    )


def build_document():
    doc = Document()
    setup_document(doc)
    add_title(doc)

    doc.add_heading("1. Аннотация", level=1)
    doc.add_paragraph(
        "Этот отчет описывает pipeline для перевода с чагатайского на "
        "английский через in-context learning (ICL). В центре эксперимента не "
        "обучение новой translation model, а выбор демонстраций: какие train "
        "examples попадут в prompt перед вызовом LLM. Сравниваются обычные "
        "retrieval baseline'ы (BM25 и dense similarity), графовые методы "
        "(graph_common, graph_ppr) и hybrid retrieval, где объединяются "
        "semantic, lexical и graph-structural signals. Отдельно фиксируется, "
        "какой prompt уходил в LLM: без этого сравнение retrieval methods "
        "неполное, потому что prompt_mode меняет не только examples, но и "
        "наличие dictionary hints."
    )
    add_note(
        doc,
        "Главная идея",
        "Граф не является моделью перевода. Он является retriever'ом: выбирает, "
        "какие train-примеры попадут в ICL prompt. Сам перевод выполняет LLM.",
    )
    add_note(
        doc,
        "Важно про размер данных",
        "10 percent runs использовались как быстрые smoke tests. Для финального "
        "сравнения добавлены 100 percent checks на том же English phrase/sentence "
        "test split. Поэтому в таблицах ниже явно указана доля данных.",
    )

    doc.add_heading("2. Датасет и постановка эксперимента", level=1)
    doc.add_paragraph(
        "Датасет представляет собой небольшой параллельный корпус на основе "
        "учебникового материала по чагатайскому языку. Основная оценка "
        "использует English target language и phrase/sentence examples, где "
        "многословность определяется со стороны Chagatai source. В старом "
        "train/test split это дает 651 train rows и 139 test rows; dev split "
        "не использовался в текущих запусках."
    )
    add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ("source_lang", "chg"),
            ("target_lang", "en"),
            ("sample_fraction", "0.1 for smoke tests; none for 100 percent checks"),
            ("only_sentences", "true"),
            ("100 percent train rows", "651"),
            ("100 percent test rows", "139"),
            ("10 percent smoke rows", "66 train / 14 test"),
            ("k", "3 for DIPMT runs; 8 for graph/dense/BM25 baselines"),
            ("temperature", "0.0"),
        ],
        widths=[2.1, 4.3],
    )

    doc.add_heading("3. Полный ICL pipeline", level=1)
    doc.add_paragraph(
        "Для каждого test-примера pipeline выбирает k=8 примеров из train, "
        "форматирует их как few-shot demonstrations, добавляет новый "
        "чагатайский input, вызывает LLM и сравнивает prediction с reference."
    )
    add_formula(
        doc,
        "test source_text -> retriever -> top-k train examples -> prompt -> "
        "LLM -> prediction -> metrics",
    )
    doc.add_paragraph(
        "Prompt template специально оставлен простым: инструкция, выбранные "
        "examples и новый source text. Retrieval scores и graph features не "
        "показываются LLM; в prompt вставляются только выбранные пары "
        "`source_text -> target_text`."
    )
    add_formula(
        doc,
        "Chagatai: <source_text>\nEnglish: <target_text>\n...\n"
        "Now translate:\nChagatai: <test_source>\nEnglish:",
    )

    doc.add_heading("4. Обычные retrieval baseline'ы", level=1)
    doc.add_heading("4.1 BM25 lexical retrieval", level=2)
    doc.add_paragraph(
        "BM25 -- это лексический baseline. Он токенизирует source texts, строит "
        "BM25 index по train examples и ранжирует примеры по token-level "
        "совпадению с query."
    )
    add_formula(doc, "score_BM25(q, s) = BM25(query_tokens, source_tokens)")

    doc.add_heading("4.2 Dense embedding retrieval", level=2)
    doc.add_paragraph(
        "Dense retrieval кодирует source texts через multilingual sentence "
        "transformer, например BGE-M3 или intfloat/multilingual-e5-base. В "
        "текущем тройном гибриде используется intfloat/multilingual-e5-base. "
        "Train examples ранжируются по cosine similarity между normalized "
        "embeddings."
    )
    add_formula(
        doc,
        "score_dense(q, s) = cosine(embedding(q), embedding(s)) = e_q^T e_s",
    )

    doc.add_heading("5. Graph-aware retrieval", level=1)
    doc.add_paragraph(
        "Graph retriever строит взвешенный двудольный граф из train set. Левая "
        "доля содержит train source examples, правая доля содержит feature nodes, "
        "извлеченные только из Chagatai source text и легких structural signals "
        "типа `type` и `length`."
    )
    add_formula(
        doc,
        "G = (S union F, E), where S = train examples and F = feature nodes",
    )

    doc.add_heading("5.1 Построение features", level=2)
    doc.add_paragraph(
        "Для каждого train source извлекаются lexical, subword-like, sequence и "
        "structural features. Эти признаки становятся feature nodes, соединенными "
        "с соответствующим source node."
    )
    add_table(
        doc,
        ["Тип feature", "Пример", "Base weight"],
        [
            ("token", "token:tuz", "3.0"),
            ("bigram", "bigram:wa tuzdur", "2.2"),
            ("suffix3", "suffix3:dur", "1.2"),
            ("suffix4", "suffix4:zdur", "0.9"),
            ("prefix3", "prefix3:tuz", "0.6"),
            ("type", "type:phrase", "1.0"),
            ("length bucket", "length:3-4", "0.8"),
        ],
        widths=[1.65, 3.4, 1.35],
    )
    doc.add_paragraph(
        "Финальный вес ребра дополнительно умножается на IDF: редкие признаки "
        "становятся информативнее частых признаков."
    )
    add_formula(
        doc,
        "idf(f) = log((N + 1) / (df(f) + 1)) + 1\nw(s, f) = base_weight(f) * idf(f)",
    )

    doc.add_heading("5.2 graph_common", level=2)
    doc.add_paragraph(
        "Стратегия graph_common превращает query и каждый train example в sparse "
        "weighted feature vector и считает cosine similarity по этим вручную "
        "построенным graph features. Это не cosine по neural embeddings, а "
        "feature-overlap cosine."
    )
    add_formula(
        doc,
        "score_common(q, s) = (v_q · v_s) / (||v_q|| ||v_s||)",
    )

    doc.add_heading("5.3 graph_ppr", level=2)
    doc.add_paragraph(
        "Стратегия graph_ppr запускает Personalized PageRank-style random walk "
        "по двудольному графу. Walk стартует из normalized query feature weights, "
        "переходит от features к train examples и обратно, а затем регулярно "
        "возвращается к исходным query features."
    )
    add_formula(
        doc,
        "P(s | f) = w(s,f) / sum_s' w(s',f)\n"
        "P(f | s) = w(s,f) / sum_f' w(s,f')\n"
        "p_S(t) = p_F(t) P(S|F)\n"
        "p_F(t+1) = (1-alpha) p_S(t) P(F|S) + alpha q\n"
        "alpha = 0.35, steps = 8",
    )
    doc.add_paragraph(
        "Финальный score -- это probability mass, назначенная train source node "
        "после walk. Это graph-structural relevance score, а не LLM confidence "
        "и не translation probability."
    )

    doc.add_heading("5.4 hybrid_graph", level=2)
    doc.add_paragraph(
        "Стратегия hybrid_graph объединяет semantic dense similarity с двумя "
        "graph scores. Это стабилизирует retrieval на маленьких sample'ах, потому "
        "что частые graph features, например очень частые токены, могут слишком "
        "сильно влиять на чисто графовый ranking."
    )
    add_formula(
        doc,
        "score_hybrid = 0.55 * score_dense + 0.30 * score_ppr + 0.15 * score_common",
    )

    doc.add_heading("5.5 graph_ppr_bm25_dense", level=2)
    doc.add_paragraph(
        "Новая стратегия проверяет другой hybrid assumption: вместо "
        "graph_common к PPR добавляется BM25. Таким образом retrieval получает "
        "три разных источника информации: dense intfloat для семантики, graph "
        "PPR для структурной близости через feature graph и BM25 для прямого "
        "лексического совпадения."
    )
    add_formula(
        doc,
        "score = 0.55 * norm(dense_intfloat)\n"
        "      + 0.30 * norm(graph_ppr)\n"
        "      + 0.15 * norm(BM25)",
    )
    doc.add_paragraph(
        "Нормализация нужна потому, что raw scores находятся в разных шкалах: "
        "cosine similarity обычно ограничен около [-1, 1], BM25 зависит от "
        "частот токенов, а PPR является probability mass на узлах train examples."
    )
    doc.add_paragraph(
        "В финальной версии structural part графа задается через `type` и "
        "`length`: `type` имеет вес 1.0, а `length` имеет вес 0.8. Эти признаки "
        "помогают отличать короткие phrase examples от длинных sentence examples."
    )

    doc.add_heading("6. Какие prompts уходили в LLM", level=1)
    doc.add_paragraph(
        "Во всех экспериментах LLM получает только текстовый prompt. Retrieval "
        "scores, graph nodes, BM25 weights и embeddings в prompt не вставляются. "
        "Они используются только до вызова LLM, чтобы выбрать demonstrations."
    )
    add_table(
        doc,
        ["Метод", "Prompt", "Что видел LLM"],
        [
            (
                "BM25 baseline",
                "standard ICL",
                "Instruction + k BM25 examples + final Chagatai input.",
            ),
            (
                "BGE-M3 dense",
                "standard ICL",
                "Same template; examples selected by embedding cosine.",
            ),
            (
                "graph_common",
                "standard ICL",
                "Same template; examples selected by graph feature cosine.",
            ),
            (
                "graph_ppr",
                "standard ICL",
                "Same template; examples selected by PPR random walk.",
            ),
            (
                "hybrid_graph",
                "standard ICL",
                "Same template; examples selected by dense + PPR + common.",
            ),
            (
                "graph_ppr_bm25_dense",
                "standard ICL",
                "Same template; examples selected by dense + PPR + BM25.",
            ),
            (
                "BM25 + DIPMT+",
                "DIPMT+",
                "BM25 examples plus dictionary hints for examples and input.",
            ),
            (
                "GIZA++/BLI + DIPMT",
                "DIPMT+",
                "Same DIPMT+ template, but lexicon is expanded/noisier.",
            ),
        ],
        widths=[1.65, 1.2, 3.55],
    )

    doc.add_heading("6.1 Standard ICL prompt", level=2)
    doc.add_paragraph(
        "Этот template использовался всеми graph methods и обычными retrieval "
        "baseline'ами, когда prompt_mode=standard."
    )
    add_formula(
        doc,
        "Task: Translate the following Chagatai texts to English.\n"
        "Output only the translation. Do not explain anything.\n\n"
        "Examples:\n"
        "Chagatai: <retrieved source 1>\n"
        "English: <retrieved target 1>\n\n"
        "...\n"
        "Now translate:\n"
        "Chagatai: <test source>\n"
        "English:",
    )
    doc.add_paragraph(
        "Для test item `şorpa göşt wa tuzdur` стратегия "
        "graph_ppr_bm25_dense поставила первым example `tuz aççıqdur -> salt is "
        "bitter`, потому что он содержит ключевой token `tuz` и связанный pattern "
        "`dur`."
    )
    add_formula(
        doc,
        "Examples:\n"
        "Chagatai: tuz aççıqdur\n"
        "English: salt is bitter\n\n"
        "Chagatai: tuz şı̇̄rı̇̄n emes\n"
        "English: salt is not sweet\n\n"
        "Now translate:\n"
        "Chagatai: şorpa göşt wa tuzdur\n"
        "English:",
    )

    doc.add_heading("6.2 DIPMT+ prompt", level=2)
    doc.add_paragraph(
        "DIPMT+ меняет prompt: к каждому retrieved example и к final input "
        "добавляются vocabulary hints из словаря. Поэтому это уже не только "
        "retrieval comparison, но и prompt augmentation."
    )
    add_formula(
        doc,
        "Task: Translate the following Chagatai texts to English.\n"
        "Use the vocabulary hints and the examples as in-context guidance.\n"
        "Output only the translation. Do not explain anything.\n\n"
        "Example 1:\n"
        "Chagatai: <retrieved source>\n"
        "Vocabulary hints:\n"
        "- \"wa\" means \"and\".\n"
        "English: <retrieved target>\n\n"
        "Now translate:\n"
        "Chagatai: şorpa göşt wa tuzdur\n"
        "Vocabulary hints:\n"
        "- \"şorpa\" means \"soup\".\n"
        "- \"göşt\" means \"meat\".\n"
        "- \"wa\" means \"and\".\n"
        "- \"tuzdur\" may contain dictionary word \"tuz\", which means \"salt\".\n"
        "- \"tuzdur\" may contain dictionary word \"dur\", "
        "which means \"is\" or \"are\".\n"
        "English:",
    )
    doc.add_paragraph(
        "В GIZA++/BLI + DIPMT варианте hints стали шумнее. Например, для тех же "
        "examples появлялись пары вроде `wa -> and/of`, `qaṣāyıd -> "
        "qasidas/knowing`, `Fārs -> short`. Это объясняет, почему chrF может "
        "остаться высоким, но BLEU и BERTScore ухудшаются."
    )
    add_formula(
        doc,
        "GIZA/BLI-style hints example:\n"
        "- \"wa\" means \"and\" or \"of\".\n"
        "- \"qānūnı\" means \"first\" or \"law\".\n"
        "- \"budur\" means \"tales\" or \"This\".",
    )

    doc.add_heading("6.3 Real test item used in examples", level=2)
    doc.add_paragraph(
        "Ниже один и тот же test item, на котором видно различие prompt content."
    )
    add_table(
        doc,
        ["Поле", "Значение"],
        [
            ("source_text", "şorpa göşt wa tuzdur"),
            ("reference", "soup is meat and salt"),
            ("type", "phrase"),
            ("length", "4"),
        ],
        widths=[1.35, 5.05],
    )

    doc.add_heading("7. Результаты", level=1)
    doc.add_paragraph(
        "Все строки ниже используют одну и ту же LLM и target language: "
        "openai/gpt-oss-120b, English. BERTScore считается через "
        "bert-base-multilingual-cased. Колонка Data used показывает, был ли это "
        "10 percent smoke test или 100 percent check."
    )
    add_table(
        doc,
        [
            "Метод",
            "Data used",
            "Prompt",
            "k",
            "BLEU",
            "chrF",
            "BERTScore",
            "Format err",
            "Empty",
        ],
        [
            (
                "BM25 baseline",
                "10%",
                "standard",
                "8",
                "6.77",
                "34.88",
                "80.98",
                "7.14",
                "0",
            ),
            (
                "BGE-M3 dense",
                "10%",
                "standard",
                "8",
                "10.29",
                "35.88",
                "79.06",
                "7.14",
                "0",
            ),
            (
                "common",
                "10%",
                "standard",
                "8",
                "11.57",
                "35.25",
                "80.76",
                "0.00",
                "0",
            ),
            ("PPR", "10%", "standard", "8", "11.75", "31.87", "79.80", "7.14", "0"),
            ("hybrid", "10%", "standard", "8", "15.02", "36.73", "81.02", "0.00", "0"),
            (
                "PPR + BM25 + E5",
                "10%",
                "standard",
                "8",
                "13.72",
                "39.41",
                "81.04",
                "7.14",
                "0",
            ),
            (
                "BM25 + DIPMT+",
                "10%",
                "DIPMT+",
                "3",
                "12.73",
                "39.44",
                "82.20",
                "7.14",
                "0",
            ),
            (
                "GIZA++/BLI + DIPMT",
                "10%",
                "DIPMT+",
                "3",
                "9.51",
                "39.57",
                "80.76",
                "7.14",
                "0",
            ),
            (
                "PPR + BM25 + E5",
                "100%",
                "standard",
                "8",
                "25.36",
                "46.24",
                "84.19",
                "6.47",
                "0",
            ),
            (
                "BM25 + DIPMT+",
                "100%",
                "DIPMT+",
                "3",
                "18.23",
                "38.58",
                "84.43",
                "5.04",
                "28",
            ),
        ],
        widths=[1.35, 0.55, 0.8, 0.35, 0.52, 0.52, 0.72, 0.62, 0.45],
    )
    doc.add_paragraph(
        "Строка BM25 + DIPMT+ отражает предыдущий результат, который был "
        "зафиксирован в terminal output до перезаписи метрик. Строка "
        "GIZA++/BLI + DIPMT отражает текущий более поздний run: его metrics были "
        "сохранены в старое имя файла, поэтому в отчете он отделен вручную."
    )
    add_note(
        doc,
        "Summary",
        "Лучший общий retrieval result сейчас у 100 percent PPR + BM25 + E5: "
        "BLEU 25.36, chrF 46.24, BERTScore 84.19, empty predictions 0. "
        "100 percent BM25 + DIPMT+ имеет самый высокий BERTScore 84.43, но "
        "28 empty predictions, поэтому его нужно трактовать осторожно."
    )

    doc.add_heading("8. Обсуждение и ограничения", level=1)
    doc.add_paragraph(
        "Графовые методы добавляют интерпретируемые retrieval signals: lexical "
        "overlap, subword-like endings, phrase length и type. Это особенно "
        "полезно в low-resource setting, где маленький корпус может содержать "
        "повторяющиеся грамматические паттерны."
    )
    doc.add_paragraph(
        "Однако текущий граф имеет вручную заданные веса. Веса hybrid strategies "
        "тоже пока эвристические: 0.55/0.30/0.15. Кроме того, DIPMT/GIZA "
        "результаты нельзя напрямую сравнивать с standard prompt как чистый "
        "retrieval-only experiment, потому что prompt содержит дополнительные "
        "словарные подсказки."
    )
    doc.add_paragraph(
        "10 percent setup оставлен только как smoke-test evidence: retrieval pool "
        "там сильно меньше полного train set. 100 percent check подтверждает, что "
        "графовый hybrid существенно улучшается при расширении candidate pool."
    )

    doc.add_heading("9. Заключение", level=1)
    doc.add_paragraph(
        "Реализованный graph-aware ICL retrieval дает прозрачный способ выбирать "
        "demonstrations через явную структуру train corpus. После проверки на "
        "100 percent split наиболее сильным и стабильным методом является "
        "PPR + BM25 + E5: он выигрывает по BLEU и chrF и не дает пустых ответов. "
        "DIPMT+ полезен как prompt-augmentation direction, но текущий 100 percent "
        "run требует дополнительной настройки из-за empty predictions."
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
